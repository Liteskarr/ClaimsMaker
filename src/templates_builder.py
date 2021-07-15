from tempfile import NamedTemporaryFile

from docx import Document as load_document
from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.table import Table
from docxtpl import DocxTemplate


def _construct_cell(cell, style: str, text: str, bold: bool = False):
    cell.paragraphs[0].add_run('')
    cell.paragraphs[0].style = style
    cell.paragraphs[0].runs[0].text = text
    cell.paragraphs[0].runs[0].bold = bold


def _construct_table(doc: Document) -> Table:
    doc.styles.add_style('unique_table_style', WD_STYLE_TYPE.PARAGRAPH)
    doc.styles['unique_table_style'].font.size = Pt(8)
    table = doc.add_table(4, 3)
    table.style = 'Table Grid'
    label_cells = table.rows[0].cells
    _construct_cell(label_cells[0], 'Normal', '{%tc for label in labels %}', bold=True)
    _construct_cell(label_cells[1], 'Normal', '{{ label }}', bold=True)
    _construct_cell(label_cells[2], 'Normal', '{%tc endfor %}', bold=True)
    start_content_cells = table.rows[1].cells
    start = start_content_cells[0].merge(start_content_cells[1]).merge(start_content_cells[2])
    _construct_cell(start, 'unique_table_style', '{%tr for row in table %}')
    content_cells = table.rows[2].cells
    _construct_cell(content_cells[0], 'unique_table_style', '{%tc for column in row %}')
    _construct_cell(content_cells[1], 'unique_table_style', '{{ column }}')
    _construct_cell(content_cells[2], 'unique_table_style', '{%tc endfor %}')
    end_content_cells = table.rows[3].cells
    end = end_content_cells[0].merge(end_content_cells[1]).merge(end_content_cells[2])
    _construct_cell(end, 'unique_table_style', '{%tr endfor %}')
    return table


KEYWORDS = {
    'АДРЕС': '{{ address }}',
    'НАЗВАНИЕ': '{{ name }}',
    'НОМЕР': '{{ number }}',
    'СУММАДЕБЕТ': '{{ sum_of_debit }}',
    'СУММАКРЕДИТ': '{{ sum_of_credit }}',
    'ДАТА': '{{ date }}',
    'ИМЯ': '{{ name }}'
}

TABLE_WORD = 'ТАБЛИЦА'


def build_template(raw_template_path: str) -> DocxTemplate:
    doc = load_document(raw_template_path)
    for p in doc.paragraphs:
        for run in p.runs:
            for keyword, replacement in KEYWORDS.items():
                run.text = run.text.replace(keyword, replacement)
    for t in doc.tables:
        for row in map(lambda z: t.row_cells(z), range(len(t.rows))):
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for keyword, replacement in KEYWORDS.items():
                            if keyword in run.text:
                                run.text = run.text.replace(keyword, replacement)
    for p in doc.paragraphs:
        if TABLE_WORD in p.text:
            p._p.addnext(_construct_table(doc)._tbl)
            p.text = ''
    temp_file = NamedTemporaryFile(delete=False)
    doc.save(temp_file.name)
    template = DocxTemplate(temp_file.name)
    temp_file.close()
    return template
